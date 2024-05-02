from collections import Counter
from docx import Document
import aspose.words as aw
from pathlib import Path
import os
import re
def cuntWord(name:str):
    # fileList = getpath(name)
    # assert os.path.isfile(fileList[0])
    documentOld = aw.Document("a.docx")
    # documentOld.save("a.docx")
    # document = Document("a.docx")
    # paragraph = document.paragraphs
    # text:list[str] = []
    # for item in paragraph:
    #     sigleList = item.text.split(" ")
    #     onlyWordList = [re.sub("[^a-zA-Z]","" ,subItem).lower() for subItem in sigleList]
    #     text.extend(onlyWordList)
    
    # result = Counter(text)
    # print(result)

def getpath(name:str):
    cwd = Path.cwd()
    dirPath = os.path.join(cwd.parent,name)
    allFile = os.listdir(dirPath)
    fileList:list[str] = []
    for name in allFile:
        fullName = os.path.join(dirPath,name)
        fileList.append(fullName)
    return fileList

if __name__ == "__main__":
    cuntWord("store")
